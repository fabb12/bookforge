"""Test della sistemazione .docx (richiede python-docx)."""
import pytest

docx = pytest.importorskip("docx")

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from bookforge.core.docx_formatter import (
    format_docx, DocxFormatRules, _heading_level, _is_caption_paragraph,
    _is_toc_paragraph, _para_has_drawing,
)


def test_detection_helpers():
    assert _heading_level("Heading 1") == 1
    assert _heading_level("Titolo 2") == 2
    assert _heading_level("Normale") is None


def test_format_docx_headings_caption_toc(tmp_path):
    doc = docx.Document()
    doc.add_heading("Capitolo 1", level=1)
    doc.add_paragraph("Corpo  con   spazi doppi.")

    # voce indice con stile TOC 1 → non deve essere formattata come corpo
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc = doc.add_paragraph("Capitolo 1 .... 5"); toc.style = doc.styles["TOC 1"]
    assert _is_toc_paragraph(toc)

    # didascalia sopra un'immagine → va spostata sotto
    cap = doc.add_paragraph("Figura 1: schema.")
    assert _is_caption_paragraph(cap)
    img = doc.add_paragraph()
    r = img.add_run()
    r._r.append(docx.oxml.OxmlElement("w:drawing"))
    assert _para_has_drawing(img)

    src = tmp_path / "in.docx"; dst = tmp_path / "out.docx"
    doc.save(str(src))
    report = format_docx(src, dst, DocxFormatRules())

    assert report.headings_normalized >= 1
    assert report.captions_formatted >= 1
    assert report.captions_moved >= 1
    assert report.toc_updated is True

    # nel risultato l'immagine precede la didascalia
    out = docx.Document(str(dst))
    order = []
    for p in out.paragraphs:
        if _para_has_drawing(p):
            order.append("IMG")
        elif p.text.strip().startswith("Figura"):
            order.append("CAP")
    assert order == ["IMG", "CAP"]

    # updateFields impostato per l'aggiornamento dell'indice
    assert out.settings.element.find(qn("w:updateFields")) is not None
