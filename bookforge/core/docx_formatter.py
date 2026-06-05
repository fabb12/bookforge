"""Formattazione, impaginazione e correzione di file Word (.docx).

Lavoriamo DIRETTAMENTE sul .docx con python-docx invece di passare per LaTeX:
così le modifiche sono chirurgiche e *lossless* — tocchiamo solo ciò che va
sistemato (titoli, immagini, formattazione del corpo, margini) e preserviamo
tutto il resto (tabelle, note, struttura).

La pipeline applica, in ordine:
  1. margini di pagina
  2. normalizzazione degli stili dei titoli (gerarchia/corpo/spaziatura)
  3. formattazione del corpo del testo (font, interlinea, giustificato, rientri)
  4. ridimensionamento delle immagini troppo larghe + centratura
  5. pulizia (paragrafi vuoti multipli, spazi doppi)
  6. correzione del testo via AI (opzionale, tramite callback)

La correzione del testo è iniettata come callable esterno per non accoppiare
questo modulo al motore degli agenti.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import Callable


# --------------------------------------------------------------------- regole
@dataclass
class DocxFormatRules:
    """Tutte le scelte di formattazione, con default tipograficamente sensati."""
    # --- corpo del testo ---
    format_body: bool = True
    body_font: str = "Times New Roman"
    body_size_pt: float = 12.0
    line_spacing: float = 1.5
    justify: bool = True
    first_line_indent_cm: float = 0.5
    space_after_pt: float = 6.0

    # --- titoli ---
    normalize_headings: bool = True
    heading_font: str = ""              # vuoto = stesso font del corpo
    heading1_size_pt: float = 18.0
    heading2_size_pt: float = 15.0
    heading3_size_pt: float = 13.0
    heading_bold: bool = True
    heading_space_before_pt: float = 14.0
    heading_space_after_pt: float = 8.0

    # --- immagini ---
    fit_images_to_page: bool = True
    max_image_width_cm: float = 0.0     # 0 = larghezza utile della pagina
    center_images: bool = True

    # --- didascalie (caption) ---
    format_captions: bool = True
    caption_below_image: bool = True    # sposta la didascalia SOTTO l'immagine
    caption_size_pt: float = 10.0
    caption_italic: bool = True
    caption_center: bool = True

    # --- indice / sommario (TOC) ---
    fix_toc: bool = True                # forza Word ad aggiornare l'indice all'apertura

    # --- pulizia ---
    remove_empty_paragraphs: bool = True
    collapse_spaces: bool = True

    # --- pagina ---
    set_margins: bool = True
    margin_cm: float = 2.5

    # --- correzione AI (gestita dal chiamante via callback) ---
    correct_text: bool = False

    def to_dict(self) -> dict:
        return asdict(self)

    @staticmethod
    def from_dict(d: dict) -> "DocxFormatRules":
        import dataclasses
        known = {f.name for f in dataclasses.fields(DocxFormatRules)}
        return DocxFormatRules(**{k: v for k, v in d.items() if k in known})


@dataclass
class FormatReport:
    """Riepilogo di cosa è stato modificato, da mostrare all'utente."""
    headings_normalized: int = 0
    body_paragraphs_formatted: int = 0
    images_resized: int = 0
    images_centered: int = 0
    captions_formatted: int = 0
    captions_moved: int = 0
    empty_paragraphs_removed: int = 0
    paragraphs_corrected: int = 0
    margins_set: bool = False
    toc_updated: bool = False
    warnings: list[str] = field(default_factory=list)
    output_path: str = ""

    def summary(self) -> str:
        lines = [
            f"Titoli normalizzati: {self.headings_normalized}",
            f"Paragrafi formattati: {self.body_paragraphs_formatted}",
            f"Immagini ridimensionate: {self.images_resized}",
            f"Immagini centrate: {self.images_centered}",
            f"Didascalie sistemate: {self.captions_formatted}",
        ]
        if self.captions_moved:
            lines.append(f"Didascalie spostate sotto l'immagine: {self.captions_moved}")
        lines.append(f"Paragrafi vuoti rimossi: {self.empty_paragraphs_removed}")
        if self.paragraphs_corrected:
            lines.append(f"Paragrafi corretti (AI): {self.paragraphs_corrected}")
        if self.margins_set:
            lines.append("Margini di pagina impostati.")
        if self.toc_updated:
            lines.append("Indice: Word lo aggiornerà all'apertura del documento.")
        if self.warnings:
            lines.append("\nAvvisi:")
            lines.extend(f"  • {w}" for w in self.warnings)
        if self.output_path:
            lines.append(f"\nFile salvato in:\n{self.output_path}")
        return "\n".join(lines)


# --------------------------------------------------------------- riconoscimenti
# Word usa "Heading N" (EN) o "Titolo N" (IT) per i livelli dei titoli.
_HEADING_PREFIXES = ("heading", "titolo")


def _heading_level(style_name: str) -> int | None:
    """Restituisce il livello (1,2,3...) se lo stile è un titolo, altrimenti None."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    for pref in _HEADING_PREFIXES:
        if name.startswith(pref):
            tail = name[len(pref):].strip()
            if tail.isdigit():
                return int(tail)
    return None


# Stili che Word usa per le voci dell'indice (EN "TOC N" / IT "Indice N" o "Sommario N").
_TOC_PREFIXES = ("toc", "indice", "sommario", "table of contents")

# Stili e parole-chiave tipici delle didascalie (EN "Caption" / IT "Didascalia").
_CAPTION_STYLE_HINTS = ("caption", "didascalia")
_CAPTION_TEXT_RE = re.compile(
    r"^\s*(fig(?:ura|ure|\.)?|tab(?:ella|le|\.)?|immagine|foto|grafico|schema|tav(?:ola|\.)?)\b"
    r"[\s\.:]*\d",
    re.IGNORECASE,
)


def _style_name(para) -> str:
    try:
        return (para.style.name if para.style else "").strip().lower()
    except Exception:  # noqa: BLE001
        return ""


def _is_toc_paragraph(para) -> bool:
    """Una voce dell'indice: stile «TOC N» / «Indice N» / «Sommario N»."""
    name = _style_name(para)
    return any(name.startswith(p) for p in _TOC_PREFIXES)


def _para_has_drawing(para) -> bool:
    return bool(para._p.findall(".//" + _drawing_tag()))


def _is_caption_paragraph(para) -> bool:
    """Riconosce una didascalia da stile o da testo (Figura/Tabella N…)."""
    name = _style_name(para)
    if any(hint in name for hint in _CAPTION_STYLE_HINTS):
        return True
    txt = para.text.strip()
    if txt and len(txt) < 320 and not _para_has_drawing(para) \
            and _CAPTION_TEXT_RE.match(txt):
        return True
    return False


# ------------------------------------------------------------------- formatter
def format_docx(
    src_path: str | Path,
    dst_path: str | Path,
    rules: DocxFormatRules | None = None,
    text_corrector: Callable[[str], str] | None = None,
    progress: Callable[[str], None] | None = None,
) -> FormatReport:
    """Apre `src_path`, applica `rules`, salva in `dst_path` e restituisce un report.

    `text_corrector`, se fornito e `rules.correct_text` è attivo, viene chiamato
    su ogni paragrafo del corpo per restituirne la versione corretta.
    """
    try:
        import docx
        from docx.shared import Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:  # pragma: no cover - dipende dall'ambiente
        raise RuntimeError(
            "Il pacchetto 'python-docx' non è installato.\n"
            "Installa con:  pip install python-docx"
        ) from e

    rules = rules or DocxFormatRules()
    report = FormatReport()

    def step(msg: str) -> None:
        if progress:
            progress(msg)

    src_path = Path(src_path)
    dst_path = Path(dst_path)
    if not src_path.exists():
        raise FileNotFoundError(f"File non trovato: {src_path}")

    step("Apertura del documento…")
    document = docx.Document(str(src_path))

    heading_font = rules.heading_font.strip() or rules.body_font
    heading_sizes = {
        1: rules.heading1_size_pt,
        2: rules.heading2_size_pt,
        3: rules.heading3_size_pt,
    }

    # --- 1. margini di pagina ---------------------------------------------
    if rules.set_margins:
        step("Impostazione dei margini…")
        for section in document.sections:
            m = Cm(rules.margin_cm)
            section.top_margin = section.bottom_margin = m
            section.left_margin = section.right_margin = m
        report.margins_set = True

    # --- aggiorna anche lo stile base, così le modifiche "tengono" ovunque -
    if rules.format_body:
        _apply_base_style(document, rules, Pt, Cm, WD_ALIGN_PARAGRAPH)

    # --- 2/3. paragrafi: titoli e corpo -----------------------------------
    step("Formattazione di titoli e corpo del testo…")
    for para in document.paragraphs:
        level = _heading_level(para.style.name if para.style else "")

        if level is not None and rules.normalize_headings:
            size = heading_sizes.get(level, rules.heading3_size_pt)
            _apply_run_format(para, heading_font, size, bold=rules.heading_bold,
                              Pt=Pt)
            pf = para.paragraph_format
            pf.space_before = Pt(rules.heading_space_before_pt)
            pf.space_after = Pt(rules.heading_space_after_pt)
            report.headings_normalized += 1
            continue

        if level is not None:
            continue  # titolo ma normalizzazione disattivata → non toccarlo

        # Non riformattare le voci dell'indice né le didascalie: hanno regole
        # proprie (altrimenti l'indice e le caption verrebbero "sporcati").
        if _is_toc_paragraph(para):
            continue
        if rules.format_captions and _is_caption_paragraph(para):
            continue

        if rules.format_body:
            # correzione AI del testo (opzionale) — preserva lo stile del paragrafo
            if rules.correct_text and text_corrector and para.text.strip():
                try:
                    corrected = text_corrector(para.text)
                    if corrected and corrected.strip() and corrected != para.text:
                        _replace_paragraph_text(para, corrected)
                        report.paragraphs_corrected += 1
                except Exception as exc:  # noqa: BLE001
                    report.warnings.append(f"Correzione saltata su un paragrafo: {exc}")

            if rules.collapse_spaces:
                _collapse_runs_spaces(para)

            _apply_run_format(para, rules.body_font, rules.body_size_pt,
                              bold=None, Pt=Pt)
            pf = para.paragraph_format
            pf.line_spacing = rules.line_spacing
            pf.space_after = Pt(rules.space_after_pt)
            if rules.first_line_indent_cm > 0:
                pf.first_line_indent = Cm(rules.first_line_indent_cm)
            if rules.justify:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            report.body_paragraphs_formatted += 1

    # --- 4. immagini -------------------------------------------------------
    if rules.fit_images_to_page or rules.center_images:
        step("Adattamento delle immagini…")
        usable_emu = _usable_width_emu(document, rules, Cm, Emu)
        for shape in document.inline_shapes:
            try:
                if rules.fit_images_to_page and shape.width and shape.width > usable_emu:
                    ratio = usable_emu / shape.width
                    shape.height = Emu(int(shape.height * ratio))
                    shape.width = Emu(int(usable_emu))
                    report.images_resized += 1
            except Exception as exc:  # noqa: BLE001
                report.warnings.append(f"Immagine non ridimensionata: {exc}")

        if rules.center_images:
            report.images_centered = _center_inline_images(document, WD_ALIGN_PARAGRAPH)

    # --- 4b. didascalie: stile + posizionamento SOTTO l'immagine ----------
    if rules.format_captions:
        step("Sistemazione delle didascalie…")
        _format_captions(document, rules, report, Pt, WD_ALIGN_PARAGRAPH)

    # --- 4c. indice/sommario: forza Word ad aggiornarlo all'apertura ------
    if rules.fix_toc:
        step("Sistemazione dell'indice…")
        report.toc_updated = _fix_toc(document)

    # --- 5. pulizia: paragrafi vuoti multipli -----------------------------
    if rules.remove_empty_paragraphs:
        step("Pulizia dei paragrafi vuoti…")
        report.empty_paragraphs_removed = _remove_consecutive_empty(document)

    # --- salvataggio -------------------------------------------------------
    step("Salvataggio…")
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dst_path))
    report.output_path = str(dst_path)
    step("Completato.")
    return report


# ------------------------------------------------------------------- helpers
def _apply_run_format(para, font_name, size_pt, bold, Pt) -> None:
    """Applica font e corpo a tutti i run; `bold=None` non tocca il grassetto."""
    runs = para.runs
    if not runs:
        return
    for run in runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        # imposta anche il font per i caratteri east-asian/complex per coerenza
        _set_run_eastasia_font(run, font_name)
        if bold is not None:
            run.font.bold = bold


def _set_run_eastasia_font(run, font_name) -> None:
    """python-docx imposta solo w:ascii; forziamo anche w:eastAsia/w:hAnsi/w:cs."""
    try:
        from docx.oxml.ns import qn
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_name)
    except Exception:  # noqa: BLE001 - best effort, non bloccante
        pass


def _replace_paragraph_text(para, new_text: str) -> None:
    """Sostituisce il testo del paragrafo conservando lo stile del primo run.

    Nota: la formattazione *inline* (grassetto/corsivo su singole parole) viene
    persa, perché la correzione AI lavora sul testo piano del paragrafo.
    """
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _collapse_runs_spaces(para) -> None:
    """Riduce sequenze di spazi multipli a uno solo, run per run."""
    import re
    for run in para.runs:
        if run.text:
            run.text = re.sub(r"[ \t]{2,}", " ", run.text)


def _usable_width_emu(document, rules, Cm, Emu) -> float:
    """Larghezza utile (pagina meno margini) in EMU, o il limite scelto dall'utente."""
    if rules.max_image_width_cm and rules.max_image_width_cm > 0:
        return float(int(Cm(rules.max_image_width_cm)))
    section = document.sections[0]
    page_w = section.page_width or Cm(21)          # default A4
    left = section.left_margin or Cm(rules.margin_cm)
    right = section.right_margin or Cm(rules.margin_cm)
    return float(int(page_w - left - right))


def _center_inline_images(document, WD_ALIGN_PARAGRAPH) -> int:
    """Centra i paragrafi che contengono (solo) un'immagine inline."""
    from docx.oxml.ns import qn
    count = 0
    for para in document.paragraphs:
        # un paragrafo "immagine" ha un drawing ma testo vuoto
        has_drawing = bool(para._p.findall(".//" + qn("w:drawing")))
        if has_drawing and not para.text.strip():
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            count += 1
    return count


def _remove_consecutive_empty(document) -> int:
    """Rimuove i paragrafi vuoti consecutivi (ne lascia al massimo uno)."""
    removed = 0
    prev_empty = False
    for para in list(document.paragraphs):
        is_empty = not para.text.strip() and not para._p.findall(
            ".//" + _drawing_tag())
        if is_empty and prev_empty:
            _delete_paragraph(para)
            removed += 1
            continue
        prev_empty = is_empty
    return removed


def _drawing_tag() -> str:
    from docx.oxml.ns import qn
    return qn("w:drawing")


def _delete_paragraph(para) -> None:
    el = para._element
    el.getparent().remove(el)
    el._p = el._element = None


# ------------------------------------------------------------------- stile base
def _find_style(document, names):
    """Cerca uno stile per nome (tollerante a IT/EN), restituisce l'oggetto o None."""
    for n in names:
        try:
            return document.styles[n]
        except KeyError:
            continue
    return None


def _apply_base_style(document, rules, Pt, Cm, WD_ALIGN_PARAGRAPH) -> None:
    """Aggiorna lo stile «Normale/Normal» così la formattazione del corpo «tiene»
    anche dove i paragrafi non hanno formattazione diretta sui run."""
    normal = _find_style(document, ("Normal", "Normale"))
    if normal is None:
        return
    try:
        normal.font.name = rules.body_font
        normal.font.size = Pt(rules.body_size_pt)
        _set_style_eastasia_font(normal, rules.body_font)
        pf = normal.paragraph_format
        pf.line_spacing = rules.line_spacing
        pf.space_after = Pt(rules.space_after_pt)
        if rules.first_line_indent_cm > 0:
            pf.first_line_indent = Cm(rules.first_line_indent_cm)
        if rules.justify:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:  # noqa: BLE001 - best effort
        pass


def _set_style_eastasia_font(style, font_name) -> None:
    try:
        from docx.oxml.ns import qn
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_name)
    except Exception:  # noqa: BLE001
        pass


# ------------------------------------------------------------------- didascalie
def _format_captions(document, rules, report, Pt, WD_ALIGN_PARAGRAPH) -> None:
    """Stila le didascalie (corpo più piccolo, corsivo, centrato) e — quando
    serve — le sposta SOTTO l'immagine a cui si riferiscono."""
    paras = document.paragraphs
    n = len(paras)
    for i, para in enumerate(paras):
        if not _is_caption_paragraph(para):
            continue

        # stile della didascalia
        for run in para.runs:
            run.font.size = Pt(rules.caption_size_pt)
            if rules.caption_italic:
                run.font.italic = True
        if rules.caption_center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # una didascalia non va rientrata come un paragrafo di corpo
        para.paragraph_format.first_line_indent = None
        report.captions_formatted += 1

        # posizionamento: se la didascalia sta SOPRA l'immagine, spostala sotto
        if rules.caption_below_image and i + 1 < n:
            nxt = paras[i + 1]
            prev = paras[i - 1] if i > 0 else None
            next_is_img = _para_has_drawing(nxt)
            prev_is_img = bool(prev) and _para_has_drawing(prev)
            if next_is_img and not prev_is_img:
                _move_paragraph_after(para, nxt)
                report.captions_moved += 1


def _move_paragraph_after(para, target) -> None:
    """Sposta il paragrafo `para` immediatamente dopo `target` nell'XML."""
    p_el = para._p
    t_el = target._p
    parent = p_el.getparent()
    if parent is not None:
        parent.remove(p_el)
    t_el.addnext(p_el)


# ------------------------------------------------------------------- indice/TOC
def _fix_toc(document) -> bool:
    """Fa sì che Word ricalcoli l'indice (numeri di pagina/voci) all'apertura.

    Non possiamo impaginare noi il documento, quindi marchiamo i campi TOC come
    «dirty» e impostiamo <w:updateFields/> in settings.xml: Word aggiornerà
    l'indice (e gli altri campi) appena apre il file.
    """
    from docx.oxml.ns import qn

    updated = False

    # 1) marca i campi TOC come "dirty" (campi semplici e complessi)
    body = document.element.body
    for fld in body.iter(qn("w:fldSimple")):
        if "TOC" in (fld.get(qn("w:instr")) or "").upper():
            fld.set(qn("w:dirty"), "true")
            updated = True
    for instr in body.iter(qn("w:instrText")):
        if instr.text and "TOC" in instr.text.upper():
            # risali al fldChar "begin" e marcalo dirty
            run = instr.getparent()
            prev = run.getprevious() if run is not None else None
            while prev is not None:
                fc = prev.find(qn("w:fldChar"))
                if fc is not None and fc.get(qn("w:fldCharType")) == "begin":
                    fc.set(qn("w:dirty"), "true")
                    updated = True
                    break
                prev = prev.getprevious()

    # 2) chiedi a Word di aggiornare tutti i campi all'apertura
    try:
        settings = document.settings.element
        upd = settings.find(qn("w:updateFields"))
        if upd is None:
            upd = settings.makeelement(qn("w:updateFields"), {})
            settings.insert(0, upd)
        upd.set(qn("w:val"), "true")
        updated = True
    except Exception:  # noqa: BLE001 - best effort
        pass

    return updated
